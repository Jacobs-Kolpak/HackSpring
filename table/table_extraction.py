from __future__ import annotations

import argparse
import csv
import json
import os
from dataclasses import dataclass
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence


SYSTEM_PROMPT = (
    "Ты ассистент по структурированию данных. "
    "Извлекай табличную информацию из неструктурированного текста. "
    "Возвращай только JSON и ничего кроме JSON."
)
TEXT_EXTENSIONS = {".txt", ".md", ".rst", ".log", ".ini", ".cfg", ".yaml", ".yml", ".xml", ".html"}
JSON_EXTENSIONS = {".json", ".jsonl"}


@dataclass(frozen=True)
class Settings:
    api_key: str = os.getenv("HACKAI_API_KEY", "")
    base_url: str = os.getenv("HACKAI_LLM_BASE_URL", "")
    model: str = os.getenv("HACKAI_LLM_MODEL", "gpt-oss-20b")
    temperature: float = float(os.getenv("TABLE_LLM_TEMPERATURE", "0.1"))
    max_tokens: int = int(os.getenv("TABLE_LLM_MAX_TOKENS", "4096"))
    timeout_sec: float = float(os.getenv("TABLE_LLM_TIMEOUT_SEC", "120"))


@dataclass
class ExtractedTable:
    title: str
    columns: List[str]
    rows: List[List[str]]

    def to_dict(self) -> Dict[str, Any]:
        return {"title": self.title, "columns": self.columns, "rows": self.rows}


def _build_user_prompt(text: str, hint: Optional[str]) -> str:
    hint_part = f"Контекст/подсказка: {hint}\n" if hint else ""
    return (
        "Преобразуй текст в таблицу. "
        "Выдели сущности и поля, сохрани смысл, не выдумывай факты.\n"
        "Требования:\n"
        "1) Ответ только в JSON-формате.\n"
        "2) Структура JSON строго:\n"
        "{\n"
        "  \"title\": \"Название таблицы\",\n"
        "  \"columns\": [\"col1\", \"col2\"],\n"
        "  \"rows\": [[\"v11\", \"v12\"], [\"v21\", \"v22\"]]\n"
        "}\n"
        "3) rows[i] должен иметь ту же длину, что и columns.\n"
        "4) Если данных мало, верни 1-2 колонки и минимально нужные строки.\n"
        "5) Пустые значения заполняй пустой строкой.\n"
        f"{hint_part}"
        f"Текст:\n{text}"
    )


def extract_table(
    text: str,
    *,
    hint: Optional[str] = None,
    model: Optional[str] = None,
    settings: Settings = Settings(),
) -> ExtractedTable:
    if not text or not text.strip():
        raise ValueError("Input text is empty")

    try:
        from openai import OpenAI
    except ImportError as exc:
        raise RuntimeError("Package 'openai' is required. Install: pip install openai") from exc

    client = OpenAI(
        api_key=settings.api_key,
        base_url=settings.base_url.rstrip("/") + "/v1",
        timeout=settings.timeout_sec,
        max_retries=2,
    )

    response = client.chat.completions.create(
        model=model or settings.model,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": _build_user_prompt(text, hint)},
        ],
        temperature=settings.temperature,
        max_tokens=settings.max_tokens,
    )
    raw = (response.choices[0].message.content or "").strip()
    payload = _parse_json_payload(raw)
    return _normalize_table(payload)


def _parse_json_payload(raw: str) -> Dict[str, Any]:
    if not raw:
        raise RuntimeError("LLM returned empty response")

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        start = raw.find("{")
        end = raw.rfind("}")
        if start == -1 or end == -1 or start >= end:
            raise RuntimeError("LLM response is not valid JSON") from None
        return json.loads(raw[start : end + 1])


def _normalize_table(payload: Dict[str, Any]) -> ExtractedTable:
    columns_raw = payload.get("columns")
    rows_raw = payload.get("rows")

    if not isinstance(columns_raw, list) or not columns_raw:
        raise RuntimeError("JSON missing non-empty 'columns' array")
    if not isinstance(rows_raw, list):
        raise RuntimeError("JSON missing 'rows' array")

    columns = [
        str(value).strip() if str(value).strip() else f"column_{index + 1}"
        for index, value in enumerate(columns_raw)
    ]
    width = len(columns)

    rows: List[List[str]] = []
    for row in rows_raw:
        if not isinstance(row, list):
            continue
        values = [str(value).strip() for value in row]
        if len(values) < width:
            values.extend([""] * (width - len(values)))
        elif len(values) > width:
            values = values[:width]
        rows.append(values)

    title = str(payload.get("title") or "Extracted Table").strip() or "Extracted Table"
    return ExtractedTable(title=title, columns=columns, rows=rows)


def table_to_csv_text(table: ExtractedTable) -> str:
    buffer = StringIO()
    writer = csv.writer(buffer)
    writer.writerow(table.columns)
    writer.writerows(table.rows)
    return buffer.getvalue()


def table_to_xlsx_bytes(table: ExtractedTable) -> bytes:
    try:
        from openpyxl import Workbook  # type: ignore
    except ImportError as exc:
        raise RuntimeError("Excel export requires openpyxl. Install: pip install openpyxl") from exc

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = table.title[:31] if table.title else "Sheet1"
    worksheet.append(table.columns)
    for row in table.rows:
        worksheet.append(row)

    stream = BytesIO()
    workbook.save(stream)
    return stream.getvalue()


def _read_input_text(text: Optional[str], input_paths: Optional[Sequence[Path]]) -> str:
    if text and text.strip():
        return text
    if input_paths:
        chunks: List[str] = []
        for path in input_paths:
            content = _read_text_from_file(path)
            if content.strip():
                chunks.append(f"### FILE: {path.name}\n{content.strip()}")
        if chunks:
            return "\n\n".join(chunks)
    raise ValueError("Provide --text or at least one --input file")


def _read_text_from_file(path: Path) -> str:
    if not path.exists() or not path.is_file():
        raise ValueError(f"Input file not found: {path}")

    suffix = path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        return _read_plain_text(path)
    if suffix in JSON_EXTENSIONS:
        return _read_json_text(path)
    if suffix == ".csv":
        return _read_csv_text(path)
    if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        return _read_xlsx_text(path)
    if suffix == ".pdf":
        return _read_pdf_text(path)
    if suffix == ".docx":
        return _read_docx_text(path)
    if suffix == ".doc":
        raise RuntimeError(
            "Формат .doc (старый Word) напрямую не поддержан. "
            "Сохраните файл как .docx и запустите снова."
        )
    return _read_plain_text(path)


def _read_plain_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            return path.read_text(encoding="cp1251")
        except UnicodeDecodeError:
            return path.read_text(encoding="latin-1")


def _read_json_text(path: Path) -> str:
    raw = _read_plain_text(path)
    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        return raw
    return json.dumps(data, ensure_ascii=False, indent=2)


def _read_csv_text(path: Path) -> str:
    lines: List[str] = []
    with path.open("r", encoding="utf-8", newline="") as handle:
        reader = csv.reader(handle)
        for row in reader:
            lines.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(lines)


def _read_xlsx_text(path: Path) -> str:
    try:
        from openpyxl import load_workbook  # type: ignore
    except ImportError as exc:
        raise RuntimeError("Для чтения .xlsx нужен openpyxl: pip install openpyxl") from exc

    workbook = load_workbook(path, read_only=True, data_only=True)
    chunks: List[str] = []
    for sheet in workbook.worksheets:
        chunks.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if cell is None else str(cell).strip() for cell in row]
            if any(values):
                chunks.append(" | ".join(values))
    workbook.close()
    return "\n".join(chunks)


def _read_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as exc:
        raise RuntimeError("Для чтения .pdf нужен pypdf: pip install pypdf") from exc

    reader = PdfReader(str(path))
    pages: List[str] = []
    for index, page in enumerate(reader.pages, start=1):
        content = (page.extract_text() or "").strip()
        if content:
            pages.append(f"[Page {index}]\n{content}")
    return "\n\n".join(pages)


def _read_docx_text(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("Для чтения .docx нужен python-docx: pip install python-docx") from exc

    document = Document(str(path))
    parts: List[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Extract structured table data from unstructured text"
    )
    parser.add_argument("--text", type=str, help="Input raw text")
    parser.add_argument(
        "--input",
        type=Path,
        nargs="+",
        help="One or more input files (.pdf, .docx, .txt, .csv, .xlsx, .json, ...)",
    )
    parser.add_argument("--hint", type=str, help="Optional extraction hint")
    parser.add_argument("--model", type=str, help="Optional model name")
    parser.add_argument(
        "--format",
        type=str,
        default="json",
        choices=["json", "csv", "xlsx"],
        help="Output format",
    )
    parser.add_argument(
        "--output",
        type=Path,
        help="Output file path for csv/xlsx/json",
    )
    return parser


def main() -> None:
    args = _build_parser().parse_args()
    text = _read_input_text(args.text, args.input)
    table = extract_table(text, hint=args.hint, model=args.model)

    if args.format == "json":
        content = json.dumps(table.to_dict(), ensure_ascii=False, indent=2)
        if args.output:
            args.output.parent.mkdir(parents=True, exist_ok=True)
            args.output.write_text(content, encoding="utf-8")
            print(str(args.output.resolve()))
        else:
            print(content)
        return

    if args.format == "csv":
        output = args.output or Path("table_output.csv")
        if output.suffix.lower() != ".csv":
            output = output.with_suffix(".csv")
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_text(table_to_csv_text(table), encoding="utf-8")
        print(str(output.resolve()))
        return

    output = args.output or Path("table_output.xlsx")
    if output.suffix.lower() != ".xlsx":
        output = output.with_suffix(".xlsx")
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_bytes(table_to_xlsx_bytes(table))
    print(str(output.resolve()))


if __name__ == "__main__":
    main()
